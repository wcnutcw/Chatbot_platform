import os
import base64
from docx import Document
import pdfplumber
import fitz  # PyMuPDF
import pandas as pd
from fastapi.responses import JSONResponse
from tempfile import TemporaryDirectory
from cleasing import cleansing
import logging

def read_docx(path):
    doc = Document(path)
    paragraphs = []
    tables = []
    images_b64 = []

    # เก็บแต่ละ paragraph แยก
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append(table_data)

    for shape in doc.inline_shapes:
        try:
            image = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = doc.part.related_parts[image]
            image_bytes = image_part.blob
            image_b64 = base64.b64encode(image_bytes).decode("utf-8")
            images_b64.append(image_b64)
        except Exception as e:
            logging.error(f"❌ Error extracting DOCX image: {e}")

    return paragraphs, tables, images_b64

def read_pdf(path):
    pages = []
    tables = []
    images_b64 = []

    # อ่านข้อความแยกแต่ละหน้า
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            try:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text.strip())
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
            except Exception as e:
                logging.error(f"❌ Error processing page {page_num}: {e}")

    # อ่านรูปจากทุกหน้า
    doc = fitz.open(path)
    for page_index in range(len(doc)):
        page = doc.load_page(page_index)
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_b64 = base64.b64encode(image_bytes).decode("utf-8")
            images_b64.append(image_b64)

    return pages, tables, images_b64

async def Up_File(upload_files):
    dfs = []
    all_tables = []
    all_images_b64 = []
    all_paragraphs = []  # สำหรับ DOCX
    all_pages = []       # สำหรับ PDF

    with TemporaryDirectory() as tmpdir:
        for upload_file in upload_files:
            filename = upload_file.filename
            file_path = os.path.join(tmpdir, filename)

            # Save file to disk temporarily
            with open(file_path, "wb") as f:
                content = await upload_file.read()
                f.write(content)

            try:
                if filename.endswith('.csv'):
                    df = pd.read_csv(file_path)
                    dfs.append(df)

                elif filename.endswith('.xlsx'):
                    excel_data = pd.read_excel(file_path, sheet_name=None)
                    for sheet in excel_data.values():
                        dfs.append(sheet)

                elif filename.endswith('.docx'):
                    paragraphs, tables, images = read_docx(file_path)
                    all_paragraphs.extend(paragraphs)
                    all_tables.extend(tables)
                    all_images_b64.extend(images)

                elif filename.endswith('.pdf'):
                    pages, tables, images = read_pdf(file_path)
                    all_pages.extend(pages)
                    all_tables.extend(tables)
                    all_images_b64.extend(images)

            except Exception as e:
                logging.error(f"❌ Error processing file {filename}: {e}")

    df_combined = pd.DataFrame()  # default กรณีไม่มีอะไรเลย

    if dfs:
        try:
            df_concat = pd.concat(dfs, ignore_index=True)
            if df_concat.empty:
                return JSONResponse(content={"error": "No valid data in CSV or Excel files"}, status_code=400)
            df_combined = cleansing(df_concat)
        except Exception as e:
            logging.error(f"❌ Error during concat or cleansing: {e}")
            return JSONResponse(content={"error": f"Error during concat or cleansing: {str(e)}"}, status_code=400)

    elif all_paragraphs:
        df_combined = pd.DataFrame({"paragraph": all_paragraphs})

    elif all_pages:
        df_combined = pd.DataFrame({"page": all_pages})

    if df_combined is None or df_combined.empty:
        return JSONResponse(content={"error": "No valid data after cleansing or file read"}, status_code=400)

    result = {
        "dataframe": df_combined.to_dict(orient="records"),
        "tables": all_tables,
        "images_b64": all_images_b64,
        "paragraphs": all_paragraphs,  # DOCX
        "pages": all_pages,            # PDF
    }

    return result
